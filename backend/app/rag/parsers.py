"""Turns a source (file path or URL) into plain text, ready for `chunker.chunk_text`.

One dispatcher, `parse()`, fans out to a per-`source_type` implementation. File
parsers (pdf/docx/epub/md/txt/html) do blocking I/O, so they run under
`asyncio.to_thread`; network parsers (web/youtube/rss) are natively async via httpx.

Anything that fetches a URL on the caller's behalf goes through
`validate_public_url` first — see its docstring for why.
"""

from __future__ import annotations

import asyncio
import re
import tempfile
from dataclasses import dataclass, field
from ipaddress import ip_address
from pathlib import Path
from typing import Any
from urllib.parse import parse_qs, urlparse

import docx
import feedparser
import httpx
import pypdf
import structlog
from bs4 import BeautifulSoup
from ebooklib import ITEM_DOCUMENT, epub
from readability import Document as ReadabilityDocument
from youtube_transcript_api import YouTubeTranscriptApi

from app.core.config import settings
from app.core.errors import UpstreamError, ValidationError

log = structlog.get_logger(__name__)

# Descriptive per courtesy to sites we fetch on a user's behalf, and so an admin
# glancing at server logs on the other end can tell who's hitting them.
_USER_AGENT = "LinguaFlowBot/1.0 (+https://linguaflow.ai; RAG content importer)"

_SOURCE_TYPES_NEEDING_PATH = {"pdf", "docx", "epub", "md", "txt", "html"}
_SOURCE_TYPES_NEEDING_URL = {"web", "youtube", "rss"}


@dataclass(slots=True)
class ParsedDoc:
    title: str
    text: str
    pages: list[tuple[int, str]] | None = None
    # Only populated for source_type="rss": one dict per feed entry
    # ({title, url, summary, published}) for the caller to fan out into
    # separate `parse("web", url=...)` calls / Document rows.
    items: list[dict[str, Any]] | None = field(default=None)


# ── SSRF guard ─────────────────────────────────────────────────────────────────


def validate_public_url(url: str) -> str:
    """Reject anything that isn't a public http(s) endpoint.

    Checked against BOTH the literal hostname in the URL and the IP(s) it
    resolves to — a DNS-rebinding attack (a public-looking hostname that
    resolves to 127.0.0.1 or a metadata-service address) defeats a
    literal-string-only check, and this importer runs on the server's behalf
    against attacker-supplied URLs (web import, RSS feeds).
    """
    parsed = urlparse(url)
    if parsed.scheme not in ("http", "https"):
        raise ValidationError(f"Unsupported URL scheme: {parsed.scheme or '(none)'}")

    host = (parsed.hostname or "").strip("[]")
    if not host:
        raise ValidationError("URL has no host.")

    host_lower = host.lower()
    if host_lower == "localhost" or host_lower.endswith(".local"):
        raise ValidationError(f"URL host is not publicly reachable: {host}")

    # Literal IP in the URL (http://127.0.0.1/...) — check before touching DNS.
    try:
        literal_ip = ip_address(host_lower)
    except ValueError:
        literal_ip = None
    if literal_ip is not None:
        _reject_if_private(literal_ip, host)
        return url  # a literal IP that passed; nothing left to resolve

    for ip in _resolve_host(host):
        _reject_if_private(ip, host)
    return url


def _resolve_host(host: str) -> list[Any]:
    """DNS-resolve `host` to every IP it maps to.

    Plain `socket.getaddrinfo` has no meaningfully-async form worth the
    complexity here — this runs once per import, not on a request hot path.
    """
    import socket

    try:
        infos = socket.getaddrinfo(host, None)
    except OSError as exc:
        raise ValidationError(f"Could not resolve host: {host}") from exc
    return [ip_address(sockaddr[0]) for *_rest, sockaddr in infos]


def _reject_if_private(ip: Any, host: str) -> None:
    # `is_private` on Python's ipaddress module already covers RFC1918 (10/8,
    # 172.16/12, 192.168/16), loopback (127/8, ::1), link-local (169.254/16,
    # fe80::/10) and unique-local IPv6 (fc00::/7) — exactly the ranges called out
    # in the spec — plus reserved/multicast/unspecified, which are worth
    # blocking too (nothing legitimate to scrape lives there).
    if ip.is_private or ip.is_loopback or ip.is_link_local or ip.is_reserved or ip.is_multicast:
        raise ValidationError(f"URL resolves to a non-public address: {host} -> {ip}")


# ── shared fetch helper (web + rss) ────────────────────────────────────────────


# Redirects are followed MANUALLY, one hop at a time, re-validating each hop.
#
# Refusing redirects outright was safe but impractical: a live RSS run lost 45 of
# 86 articles because news sites redirect constantly (http→https, canonical URLs,
# tracking-param normalisation). Handing `follow_redirects=True` to httpx is the
# other extreme — the SSRF guard would only ever inspect the first hop, so a public
# URL could bounce to 169.254.169.254.
#
# Following manually and running `validate_public_url` on EVERY hop keeps the guard
# effective while accepting ordinary web redirects.
_MAX_REDIRECTS = 5


async def _fetch_bytes(url: str) -> bytes:
    limit = settings.max_upload_bytes
    current = url

    async with httpx.AsyncClient(
        timeout=20, headers={"User-Agent": _USER_AGENT}, follow_redirects=False
    ) as client:
        for hop in range(_MAX_REDIRECTS + 1):
            # Every hop is re-validated, including the first.
            validate_public_url(current)
            try:
                async with client.stream("GET", current) as response:
                    if response.status_code in (301, 302, 303, 307, 308):
                        location = response.headers.get("location")
                        if not location:
                            raise UpstreamError(
                                f"{current} returned a redirect with no Location header."
                            )
                        if hop >= _MAX_REDIRECTS:
                            raise UpstreamError(
                                f"{url} exceeded {_MAX_REDIRECTS} redirects; giving up."
                            )
                        # Relative Locations are legal, so resolve against the
                        # current URL before the next validation pass.
                        current = str(httpx.URL(current).join(location))
                        continue

                    response.raise_for_status()
                    body = bytearray()
                    async for chunk in response.aiter_bytes():
                        body.extend(chunk)
                        if len(body) > limit:
                            raise ValidationError(
                                f"Response from {current} exceeds the {limit}-byte import limit."
                            )
                    return bytes(body)
            except httpx.HTTPStatusError as exc:
                raise UpstreamError(
                    f"{current} returned {exc.response.status_code}."
                ) from exc
            except httpx.HTTPError as exc:
                raise UpstreamError(f"Could not fetch {current}: {exc}") from exc

    raise UpstreamError(f"{url} exceeded {_MAX_REDIRECTS} redirects; giving up.")


async def _fetch_text(url: str) -> str:
    return (await _fetch_bytes(url)).decode("utf-8", errors="replace")


# ── HTML helpers ────────────────────────────────────────────────────────────────


def _html_to_text(html_or_soup: str | BeautifulSoup) -> str:
    soup = (
        html_or_soup
        if isinstance(html_or_soup, BeautifulSoup)
        else BeautifulSoup(html_or_soup, "lxml")
    )
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    lines = [ln.strip() for ln in soup.get_text("\n").splitlines()]
    return "\n".join(ln for ln in lines if ln)


def _title_from_path(path: str) -> str:
    return Path(path).stem.replace("_", " ").replace("-", " ").strip() or "Untitled"


_MD_HEADING_RE = re.compile(r"^#\s+(.+)$", re.MULTILINE)


# ── File parsers (blocking — run via asyncio.to_thread) ─────────────────────────


def _parse_pdf_sync(path: str) -> ParsedDoc:
    reader = pypdf.PdfReader(path)
    pages: list[tuple[int, str]] = []
    for i, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            pages.append((i, text))

    title = None
    try:
        if reader.metadata and reader.metadata.title:
            title = str(reader.metadata.title).strip() or None
    except Exception:  # pragma: no cover — corrupt metadata shouldn't fail the import
        pass

    return ParsedDoc(
        title=title or _title_from_path(path),
        text="\n\n".join(t for _, t in pages),
        pages=pages or None,
    )


def _parse_docx_sync(path: str) -> ParsedDoc:
    document = docx.Document(path)
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    title = paragraphs[0] if paragraphs else _title_from_path(path)
    return ParsedDoc(title=title, text="\n\n".join(paragraphs), pages=None)


def _parse_epub_sync(path: str) -> ParsedDoc:
    book = epub.read_epub(path)

    title = None
    try:
        meta = book.get_metadata("DC", "title")
        if meta:
            title = str(meta[0][0]).strip() or None
    except Exception:  # pragma: no cover — malformed OPF metadata
        pass

    parts: list[str] = []
    for item in book.get_items():
        if item.get_type() == ITEM_DOCUMENT:
            text = _html_to_text(item.get_content().decode("utf-8", errors="replace"))
            if text:
                parts.append(text)

    return ParsedDoc(title=title or _title_from_path(path), text="\n\n".join(parts), pages=None)


def _parse_text_file_sync(path: str) -> ParsedDoc:
    content = Path(path).read_text(encoding="utf-8", errors="replace")
    heading = _MD_HEADING_RE.search(content)
    title = heading.group(1).strip() if heading else _title_from_path(path)
    return ParsedDoc(title=title, text=content, pages=None)


def _parse_html_file_sync(path: str) -> ParsedDoc:
    raw = Path(path).read_text(encoding="utf-8", errors="replace")
    soup = BeautifulSoup(raw, "lxml")
    title = (
        soup.title.string.strip()
        if soup.title and soup.title.string and soup.title.string.strip()
        else _title_from_path(path)
    )
    return ParsedDoc(title=title, text=_html_to_text(soup), pages=None)


# ── Network parsers ──────────────────────────────────────────────────────────────


async def _parse_web(url: str) -> ParsedDoc:
    # Google Drive links are an application, not a page — handle them by
    # rewriting to the direct-download endpoint (or failing with instructions)
    # instead of ingesting the useless JavaScript shell.
    if is_google_drive_url(url):
        return await _parse_google_drive(url)
    html = await _fetch_text(url)
    # readability strips nav/ads/sidebars — the difference between a clean lesson
    # excerpt and a chunk full of "Subscribe to our newsletter".
    doc = ReadabilityDocument(html)
    # `short_title()` strips the site's boilerplate suffix — readability's
    # `title()` yields "Auf dem Weihnachtsmarkt (A1) | MeloLingua", which is a
    # tab caption, not a story heading. Fall back through the full title, then
    # the URL, so a page with no usable title still ingests.
    title = (doc.short_title() or "").strip() or (doc.title() or "").strip() or url
    text = _html_to_text(doc.summary(html_partial=True))
    return ParsedDoc(title=title, text=text, pages=None)


# ── Google Drive ──────────────────────────────────────────────────────────────
#
# A pasted Drive link is almost never the document itself. drive.google.com
# serves a JavaScript application: a FOLDER link renders its file listing only
# after scripts run (a plain fetch sees the empty shell, so ingesting one gave
# the learner a "document" that was just the link), and a FILE link renders a
# preview page around the file, not the file. The only fetchable form is the
# direct-download endpoint, so file links are rewritten to it and the payload
# is sniffed (pdf / docx / epub / text); folder links are rejected with
# instructions, because silently ingesting nothing is worse than an error.

_DRIVE_HOSTS = {"drive.google.com", "docs.google.com", "drive.usercontent.google.com"}
_DRIVE_FOLDER_RE = re.compile(r"/drive/(?:u/\d+/)?folders/[\w-]+")
_DRIVE_FILE_RE = re.compile(r"/file/d/([\w-]+)")
_GOOGLE_DOC_RE = re.compile(r"/document/d/([\w-]+)")


def is_google_drive_url(url: str) -> bool:
    return (urlparse(url).hostname or "").lower() in _DRIVE_HOSTS


def _drive_file_id(url: str) -> str | None:
    parsed = urlparse(url)
    match = _DRIVE_FILE_RE.search(parsed.path)
    if match:
        return match.group(1)
    if parsed.path in ("/open", "/uc"):
        ids = parse_qs(parsed.query).get("id")
        if ids and ids[0].strip():
            return ids[0].strip()
    return None


def _parse_drive_payload_sync(file_id: str, payload: bytes) -> ParsedDoc:
    """Sniff a Drive download and reuse the matching file parser.

    Content-type sniffing (magic bytes) rather than trusting an extension we
    never saw — the download URL carries no filename.
    """
    if payload.startswith(b"%PDF"):
        suffix, parsers = ".pdf", [_parse_pdf_sync]
    elif payload.startswith(b"PK\x03\x04"):
        # docx and epub are both zip containers; try both before giving up.
        suffix, parsers = ".docx", [_parse_docx_sync, _parse_epub_sync]
    else:
        head = payload[:2048].lstrip().lower()
        if head.startswith((b"<!doctype", b"<html")) or b"<html" in head:
            # An HTML body from the download endpoint is never the file. It is
            # Google's sign-in page (file not shared publicly) or the
            # can't-scan-for-viruses interstitial (file too large to fetch
            # without a confirmation cookie).
            raise UpstreamError(
                "Google Drive returned a web page instead of the file. Make sure "
                "the file is shared as 'Anyone with the link' and is small enough "
                "for a direct download, or download it and upload it here instead."
            )
        text = payload.decode("utf-8", errors="replace")
        first_line = next((ln.strip() for ln in text.splitlines() if ln.strip()), "")
        return ParsedDoc(
            title=first_line[:120] or f"Google Drive file {file_id}",
            text=text,
            pages=None,
        )

    errors: list[str] = []
    for attempt in parsers:
        # A real filename gives the parsers a sane title fallback; the id keeps
        # concurrent imports of different files from colliding.
        with tempfile.TemporaryDirectory() as tmp_dir:
            path = str(Path(tmp_dir) / f"google-drive-{file_id}{suffix}")
            Path(path).write_bytes(payload)
            try:
                return attempt(path)
            except Exception as exc:  # noqa: BLE001 — collected and re-raised below
                errors.append(f"{attempt.__name__}: {exc}")
            # epub needs its own extension for ebooklib's zip handling
            suffix = ".epub"
    raise ValidationError(
        "The Google Drive file is not a format this importer can read "
        f"(supported: PDF, DOCX, EPUB, plain text). Details: {'; '.join(errors)}"
    )


async def _parse_google_drive(url: str) -> ParsedDoc:
    parsed = urlparse(url)

    if _DRIVE_FOLDER_RE.search(parsed.path):
        raise ValidationError(
            "This is a Google Drive FOLDER link — a folder page is a JavaScript "
            "app with no document text to import. Import each file's own share "
            "link instead (right-click the file in Drive → Share → 'Anyone with "
            "the link' → copy link), or download the files and upload them here."
        )

    gdoc = _GOOGLE_DOC_RE.search(parsed.path)
    if gdoc:
        # Google Docs have a first-party plain-text export — far cleaner than
        # scraping the editor page.
        text = await _fetch_text(
            f"https://docs.google.com/document/d/{gdoc.group(1)}/export?format=txt"
        )
        stripped = text.strip()
        if not stripped or stripped.lstrip().lower().startswith(("<!doctype", "<html")):
            raise UpstreamError(
                "Google Docs did not return the document text. Make sure the doc "
                "is shared as 'Anyone with the link'."
            )
        first_line = next(ln.strip() for ln in stripped.splitlines() if ln.strip())
        return ParsedDoc(title=first_line[:120], text=stripped, pages=None)

    file_id = _drive_file_id(url)
    if not file_id:
        raise ValidationError(
            "Unrecognised Google Drive link. Use a single file's share link "
            "(…/file/d/<id>/view) or a Google Docs document link."
        )

    payload = await _fetch_bytes(f"https://drive.google.com/uc?export=download&id={file_id}")
    return await asyncio.to_thread(_parse_drive_payload_sync, file_id, payload)


_YOUTUBE_HOSTS = {"youtube.com", "www.youtube.com", "m.youtube.com", "youtube-nocookie.com"}
_YOUTUBE_ID_RE = re.compile(r"^[A-Za-z0-9_-]{11}$")


def extract_youtube_video_id(url: str) -> str:
    """Pull the 11-char video id out of any of YouTube's URL shapes.

    Handles ``watch?v=``, ``youtu.be/<id>``, ``/embed/<id>``, ``/shorts/<id>``,
    ``/live/<id>``, and a bare id passed straight through.
    """
    parsed = urlparse(url)
    host = (parsed.hostname or "").lower()

    if host == "youtu.be":
        candidate = parsed.path.lstrip("/").split("/")[0]
        if _YOUTUBE_ID_RE.match(candidate):
            return candidate

    if host in _YOUTUBE_HOSTS or host.endswith(".youtube.com"):
        if parsed.path == "/watch":
            candidate = (parse_qs(parsed.query).get("v") or [""])[0]
            if _YOUTUBE_ID_RE.match(candidate):
                return candidate
        for prefix in ("/embed/", "/shorts/", "/live/", "/v/"):
            if parsed.path.startswith(prefix):
                candidate = parsed.path[len(prefix):].split("/")[0]
                if _YOUTUBE_ID_RE.match(candidate):
                    return candidate

    if _YOUTUBE_ID_RE.match(url):  # a bare id, passed directly
        return url

    raise ValidationError(f"Could not extract a YouTube video id from: {url}")


async def _parse_youtube(url: str) -> ParsedDoc:
    video_id = extract_youtube_video_id(url)
    try:
        # German content first — this is a German-learning tutor — falling back
        # to English so an import doesn't hard-fail just because a video only
        # has auto-generated English captions.
        transcript = await asyncio.to_thread(
            YouTubeTranscriptApi().fetch, video_id, languages=("de", "en")
        )
    except Exception as exc:
        raise UpstreamError(
            f"Could not fetch a transcript for YouTube video {video_id}: {exc}"
        ) from exc

    text = " ".join(s.text.strip() for s in transcript if s.text and s.text.strip())
    return ParsedDoc(title=f"YouTube video {video_id}", text=text, pages=None)


async def _parse_rss(url: str) -> ParsedDoc:
    raw = await _fetch_bytes(url)
    feed = feedparser.parse(raw)

    items = [
        {
            "title": entry.get("title", ""),
            "url": entry.get("link", ""),
            "summary": entry.get("summary", entry.get("description", "")),
            "published": entry.get("published", entry.get("updated", "")),
        }
        for entry in feed.entries
    ]
    feed_title = (getattr(feed, "feed", {}) or {}).get("title") or url
    # No `text` — the caller (ingestion service) fans `items` out into individual
    # `parse("web", url=item["url"])` calls, one Document per entry.
    return ParsedDoc(title=feed_title, text="", pages=None, items=items)


# ── Dispatcher ───────────────────────────────────────────────────────────────────


async def parse(source_type: str, *, path: str | None = None, url: str | None = None) -> ParsedDoc:
    kind = (source_type or "").strip().lower()

    if kind in _SOURCE_TYPES_NEEDING_PATH and not path:
        raise ValidationError(f"source_type={kind!r} requires a `path`.")
    if kind in _SOURCE_TYPES_NEEDING_URL and not url:
        raise ValidationError(f"source_type={kind!r} requires a `url`.")

    if kind == "pdf":
        return await asyncio.to_thread(_parse_pdf_sync, path)  # type: ignore[arg-type]
    if kind == "docx":
        return await asyncio.to_thread(_parse_docx_sync, path)  # type: ignore[arg-type]
    if kind == "epub":
        return await asyncio.to_thread(_parse_epub_sync, path)  # type: ignore[arg-type]
    if kind in ("md", "txt"):
        return await asyncio.to_thread(_parse_text_file_sync, path)  # type: ignore[arg-type]
    if kind == "html":
        return await asyncio.to_thread(_parse_html_file_sync, path)  # type: ignore[arg-type]
    if kind == "web":
        return await _parse_web(url)  # type: ignore[arg-type]
    if kind == "youtube":
        return await _parse_youtube(url)  # type: ignore[arg-type]
    if kind == "rss":
        return await _parse_rss(url)  # type: ignore[arg-type]

    raise ValidationError(f"Unsupported source_type: {source_type!r}")
